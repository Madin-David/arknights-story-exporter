#!/usr/bin/env python3
"""
parse_test_to_docx.py

将 `test.txt` （游戏脚本样式）解析并按照《排版技术说明.md》的规范导出为 Word 文档（.docx）。

用法:
    python parse_test_to_docx.py [输入文件] [输出文件]

示例:
    python parse_test_to_docx.py test.txt formatted_test.docx

依赖: python-docx
"""
import re
import sys
import io
import hashlib
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests

# Import modular parser system
from parsers.registry import ParserRegistry
from parsers.control_parser import ControlParser
from parsers.image_parser import ImageParser
from parsers.decision_parser import DecisionParser
from parsers.predicate_parser import PredicateParser
from parsers.scene_parser import SceneParser
from parsers.subtitle_parser import SubtitleParser
from parsers.dialogue_parser import DialogueParser
from parsers.sound_parser import SoundParser
from parsers.narration_parser import NarrationParser
from config.loader import ParserConfig


def _set_font(run, font_name: str, size: float) -> None:
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def _set_paragraph_format(p, line_spacing: float = 1.5, space_before: float = 0, 
                          space_after: float = 0, first_line_indent: Optional[float] = None) -> None:
    """设置段落格式"""
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)


def create_document(page_size='A4', margin_size='normal') -> Document:
    """创建并初始化文档

    Args:
        page_size: 页面大小，'A4' 或 'Letter'，默认 'A4'
        margin_size: 页边距大小，'narrow'(窄), 'normal'(标准), 'wide'(宽)，默认 'normal'
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]

    # 设置页面大小
    if page_size == 'A4':
        section.page_height = Inches(11.69)  # A4: 297mm
        section.page_width = Inches(8.27)    # A4: 210mm
    elif page_size == 'Letter':
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)

    # 设置页边距
    margin_presets = {
        'narrow': Inches(0.5),   # 窄边距：1.27cm
        'normal': Inches(0.79),  # 标准边距：2cm
        'wide': Inches(1.0)      # 宽边距：2.54cm
    }
    margin = margin_presets.get(margin_size, Inches(0.79))
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    # 默认中文字体设置
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 默认段落格式：行距1.5倍，左对齐，段前段后0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    return doc


def add_page_numbers(doc: Document, position='bottom', alignment='center'):
    """为文档添加页码

    Args:
        doc: Document 对象
        position: 页码位置，'bottom'(底部) 或 'top'(顶部)
        alignment: 对齐方式，'left'(左对齐), 'center'(居中), 'right'(右对齐)
    """
    for section in doc.sections:
        # 获取页脚（底部）或页眉（顶部）
        if position == 'bottom':
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        else:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # 设置对齐方式
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加页码字段
        run = paragraph.add_run()

        # 创建页码字段的 XML 元素
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

        # 设置页码字体
        run.font.size = Pt(10)
        run.font.name = '宋体'


def add_main_title(doc: Document, text: str) -> None:
    """添加大标题（如"反常光谱"）：左对齐、大字号、粗体"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.bold = True
    _set_font(run, '黑体', 22)
    _set_paragraph_format(p)


def add_scene_title(doc: Document, text: str) -> None:
    """添加场景标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.bold = True
    _set_font(run, '黑体', 14)
    _set_paragraph_format(p)


def add_scene_timestamp(doc: Document, text: str) -> None:
    """添加场景时间戳"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.bold = True
    _set_font(run, '黑体', 12)
    _set_paragraph_format(p)


def add_dialogue(doc: Document, character: str, text: str) -> None:
    """添加角色对话：角色名和冒号直接连接，冒号后无空格
    角色名：宋体 8.5字号 粗体
    对话内容：宋体 8.5字号
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    char_run = p.add_run(f"{character}:")
    char_run.font.bold = True
    _set_font(char_run, '宋体', 8.5)
    
    text_run = p.add_run(text)
    _set_font(text_run, '宋体', 8.5)
    _set_paragraph_format(p)


def add_narration(doc: Document, text: str) -> None:
    """添加旁白：楷体 8.5字号，首行缩进"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, '楷体', 8.5)
    _set_paragraph_format(p, first_line_indent=0.28)


def add_image_reference(doc: Document, image_index: int):
    """添加图片索引引用（正文中）

    Args:
        image_index: 图片序号，如 1, 2, 3...
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"[图片: 图片{image_index}]")
    run.font.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
    _set_font(run, '宋体', 8.5)
    _set_paragraph_format(p)
    return run


def add_decision(doc: Document, options: list[str]) -> None:
    """添加选择支标题和选项列表

    注意：选择支列表不会显示在文档中，仅用于内部状态管理

    Args:
        options: 选项文本列表，如 ["选项1", "选项2"]
    """
    # 不输出任何内容到文档，选择支信息通过分支标题体现
    pass


def add_predicate_header(doc: Document, references: str, options_map: dict) -> None:
    """添加分支标题

    Args:
        references: 引用的选项编号，如 "1" 或 "1;2"
        options_map: 选项编号到文本的映射，如 {"1": "选项1", "2": "选项2"}
    """
    refs = references.split(';')

    # 构建分支标题文本（简洁样式，不占用整行）
    if len(refs) == 1:
        # 单个选项分支
        option_text = options_map.get(refs[0], f"选项{refs[0]}")
        title = f"[→ {option_text}]"
    elif len(refs) == len(options_map):
        # 汇合分支
        title = "[→ 汇合]"
    else:
        # 多个选项的共同分支
        option_texts = [options_map.get(r, f"选项{r}") for r in refs]
        title = f"[→ {' & '.join(option_texts)}]"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 255, 255)  # 青色
    _set_font(run, '宋体', 8.5)
    _set_paragraph_format(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def _add_sound_effect_paragraph(doc: Document, text: str) -> bool:
    """添加音效段落的内部辅助函数"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"<{text}>")
    _set_font(run, '楷体', 8.5)
    _set_paragraph_format(p)
    return True


def add_sound_effect(doc: Document, text: str) -> bool:
    """尝试将音效/音乐文本写入文档。

    规则: 如果 `text` 中包含中文字符或明显的描述性文本，则写入（如 `<雷声>`）。
    如果是资源 id（以 `$` 开头或仅包含 ASCII/下划线/连字符），则**不写入**并返回 False。
    返回 True 表示写入了 doc，False 表示已跳过（但此类行视为已识别，不会再计为未识别行）。
    """
    # 如果包含中文字符则视为描述性音效，写入
    if re.search(r'[\u4e00-\u9fff]', text):
        return _add_sound_effect_paragraph(doc, text)

    # 如果以 $ 开头或只包含 ASCII/下划线/连字符/斜杠/美元符号, 则认为是资源 id，跳过输出
    if re.match(r'^\$?[A-Za-z0-9_\-/\$]+$', text):
        return False

    # 其它情况较可能是描述性，写入
    return _add_sound_effect_paragraph(doc, text)


class DocumentAssembler:
    """Assembly helper that allows multiple parse_text calls and saves only on demand.

    Usage:
        asm = DocumentAssembler(spacer_lines=2, page_size='A4', margin_size='narrow')
        asm.parse_text(text1, title='第一章')
        asm.parse_text(text2, title='第二章')
        asm.save('out.docx')
    """

    def __init__(self, config: ParserConfig = None):
        """初始化文档组装器

        Args:
            config: ParserConfig 对象（可选，使用默认配置）
        """
        # Load configuration
        self.config = config or ParserConfig()

        # Get configuration values
        self.spacer_lines = self.config.get('document.spacer_lines', 2)
        page_size_val = self.config.get('document.page_size', 'A4')
        margin_size_val = self.config.get('document.margin_size', 'narrow')
        self.add_page_numbers = self.config.get('document.add_page_numbers', True)

        # Create document
        self.doc = create_document(page_size=page_size_val, margin_size=margin_size_val)

        # Initialize parser registry
        self.registry = ParserRegistry()
        self._setup_parsers()

        # Set context
        self.registry.context.doc = self.doc
        self.registry.context.assembler = self

        # Internal state
        self.skipped_lines = []
        self._first_section = True
        self._has_main_title = False

        # 选择支相关状态
        self.current_decision_options = {}
        self.current_predicate = None

        # 图片相关状态
        self.image_map = {}
        self.images_to_append = []
        self.image_counter = 0
        self.image_reference_runs = {}

    def _setup_parsers(self):
        """Setup parsers based on configuration"""
        parser_configs = self.config.get('parsers', {})

        # Map parser names to classes
        parser_classes = {
            'control': ControlParser,
            'image': ImageParser,
            'decision': DecisionParser,
            'predicate': PredicateParser,
            'scene': SceneParser,
            'subtitle': SubtitleParser,
            'dialogue': DialogueParser,
            'sound': SoundParser,
            'narration': NarrationParser,
        }

        # Register parsers
        for name, parser_class in parser_classes.items():
            cfg = parser_configs.get(name, {})
            enabled = cfg.get('enabled', True)
            priority = cfg.get('priority', 50)

            # Create parser with config
            if name == 'sound':
                skip_resource_ids = cfg.get('skip_resource_ids', True)
                parser = parser_class(enabled=enabled, priority=priority,
                                    skip_resource_ids=skip_resource_ids)
            else:
                parser = parser_class(enabled=enabled, priority=priority)

            self.registry.register(parser)

    def enable_parser(self, parser_name: str):
        """Enable a specific parser by name

        Args:
            parser_name: Name of the parser ('image', 'decision', etc.)
        """
        parser_classes = {
            'control': ControlParser,
            'image': ImageParser,
            'decision': DecisionParser,
            'predicate': PredicateParser,
            'scene': SceneParser,
            'subtitle': SubtitleParser,
            'dialogue': DialogueParser,
            'sound': SoundParser,
            'narration': NarrationParser,
        }
        parser_class = parser_classes.get(parser_name)
        if parser_class:
            self.registry.enable_parser(parser_class)

    def disable_parser(self, parser_name: str):
        """Disable a specific parser by name

        Args:
            parser_name: Name of the parser ('image', 'decision', etc.)
        """
        parser_classes = {
            'control': ControlParser,
            'image': ImageParser,
            'decision': DecisionParser,
            'predicate': PredicateParser,
            'scene': SceneParser,
            'subtitle': SubtitleParser,
            'dialogue': DialogueParser,
            'sound': SoundParser,
            'narration': NarrationParser,
        }
        parser_class = parser_classes.get(parser_name)
        if parser_class:
            self.registry.disable_parser(parser_class)

    def add_blank_lines(self, n: int):
        for _ in range(max(0, int(n))):
            self.doc.add_paragraph('')
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def add_main_title(self, text: str):
        """添加大标题（如"反常光谱"）"""
        if not text or self._has_main_title:
            return
        add_main_title(self.doc, text)
        self._has_main_title = True

    def add_title(self, title: str, character: str = None):
        """Add left-aligned, bold title. If character provided, prepend it without spaces."""
        if not title:
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        display_text = f"{character}{title}" if character else title
        run = p.add_run(display_text)
        run.font.bold = True
        _set_font(run, '黑体', 18)
        _set_paragraph_format(p)

    def parse_lines(self, lines, title: str = None, character: str = None, spacer_lines: int = None, image_map: dict = None):
        """Parse iterable of lines and append to internal document.

        lines: iterable of strings
        title: optional title to insert (centered, bold, large)
        character: optional character name to prepend to title (no spaces between)
        spacer_lines: overrides default spacer before this section
        image_map: 图片映射 {image_id: image_url}
        Returns the list of skipped lines accumulated so far.
        """
        # 设置图片映射
        if image_map:
            self.image_map = image_map

        # Initialize parsers
        self.registry.initialize_all()

        if title:
            n = spacer_lines if spacer_lines is not None else self.spacer_lines
            # 如果这是第一个秘录/章节, 不插入前置空行
            if not getattr(self, '_first_section', True):
                if n and n > 0:
                    self.add_blank_lines(n)
            self.add_title(title, character=character)
            # 标记已添加首个章节
            self._first_section = False

        # Parse lines using the registry
        for raw in lines:
            handled = self.registry.parse_line(raw)
            if not handled:
                self.skipped_lines.append(raw.rstrip('\n'))

        # Finalize parsers
        self.registry.finalize_all()

        return list(self.skipped_lines)

    def parse_text(self, text: str, title: str = None, character: str = None, spacer_lines: int = None, main_title: str = None, image_map: dict = None):
        """Parse text string and append to internal document.

        character: optional character name to prepend to title (no spaces between)
        main_title: optional main title (e.g., "反常光谱") to add at the beginning
        image_map: 图片映射 {image_id: image_url}
        """
        if main_title:
            self.add_main_title(main_title)
        return self.parse_lines(text.splitlines(), title=title, character=character, spacer_lines=spacer_lines, image_map=image_map)

    def append_images(self):
        """在文档末尾追加所有图片"""
        if not self.images_to_append:
            return
        hash_to_final_index = {}
        index_map = {}
        output_entries = []
        final_counter = 0

        for image_index, image_url in self.images_to_append:
            try:
                response = requests.get(image_url, timeout=10)
                response.raise_for_status()
                content = response.content
                image_hash = hashlib.sha256(content).hexdigest()

                if image_hash in hash_to_final_index:
                    index_map[image_index] = hash_to_final_index[image_hash]
                    continue

                final_counter += 1
                image_stream = io.BytesIO(content)
                output_entries.append({
                    'type': 'image',
                    'final_index': final_counter,
                    'image_stream': image_stream,
                    'image_url': image_url,
                })
                hash_to_final_index[image_hash] = final_counter
                index_map[image_index] = final_counter
            except Exception as e:
                final_counter += 1
                output_entries.append({
                    'type': 'error',
                    'final_index': final_counter,
                    'image_url': image_url,
                    'error': str(e),
                })
                index_map[image_index] = final_counter

        if not output_entries:
            return

        # 更新正文中的图片引用编号
        for original_index, run in self.image_reference_runs.items():
            final_index = index_map.get(original_index)
            if final_index:
                run.text = f"[图片: 图片{final_index}]"

        # 添加分页符
        self.add_page_break()

        # 添加图片部分标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━━━ 图片 ━━━")
        run.font.bold = True
        _set_font(run, '黑体', 14)
        _set_paragraph_format(p)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

        # 添加每张去重后的图片或错误信息
        for entry in output_entries:
            if entry['type'] == 'image':
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                entry['image_stream'].seek(0)
                run.add_picture(entry['image_stream'], width=Inches(5.5))
                _set_paragraph_format(p)

                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"图片{entry['final_index']}")
                run.font.color.rgb = RGBColor(128, 128, 128)
                _set_font(run, '宋体', 9)
                _set_paragraph_format(p)
                p.paragraph_format.space_after = Pt(18)
            else:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                error_suffix = entry.get('error')
                suffix = f" ({error_suffix})" if error_suffix else ""
                run = p.add_run(f"[图片{entry['final_index']} 加载失败: {entry['image_url']}]"
                                f"{suffix}")
                run.font.color.rgb = RGBColor(255, 0, 0)
                _set_font(run, '宋体', 9)
                _set_paragraph_format(p)
                p.paragraph_format.space_after = Pt(18)

    def save(self, outpath: str):
        # 保存前追加图片
        self.append_images()

        # 保存前添加页码
        if self.add_page_numbers:
            add_page_numbers(self.doc, position='bottom', alignment='center')

        self.doc.save(outpath)
        return list(self.skipped_lines)


def parse_line(line: str, doc: Document, assembler=None) -> bool:
    """解析一行并将可识别内容写入 doc。

    返回 True 表示该行被识别并处理，返回 False 表示该行未被识别（将被跳过）。
    assembler: DocumentAssembler 实例，用于维护选择支和图片状态
    """
    line = line.strip()
    if not line:
        return False

    # Image: 图片
    # [Image(image="27_i01")]
    m = re.search(r'Image\(image\s*=\s*"([^"]+)"', line, re.IGNORECASE)
    if m and assembler:
        image_id = m.group(1).strip()
        # 查找图片 URL
        if image_id in assembler.image_map:
            assembler.image_counter += 1
            image_url = assembler.image_map[image_id]
            # 在正文中添加图片索引
            run = add_image_reference(doc, assembler.image_counter)
            assembler.image_reference_runs[assembler.image_counter] = run
            # 记录待追加的图片
            assembler.images_to_append.append((assembler.image_counter, image_url))
        return True

    # Decision: 选择支
    # [Decision(options="选项1;选项2;...", values="1;2;...")]
    m = re.search(r'Decision\(options\s*=\s*"([^"]+)".*?values\s*=\s*"([^"]+)"', line, re.IGNORECASE)
    if m:
        options_str = m.group(1)
        values_str = m.group(2)
        options = options_str.split(';')
        values = values_str.split(';')

        # 保存选择支状态
        if assembler:
            assembler.current_decision_options = {v.strip(): o.strip() for v, o in zip(values, options)}
            assembler.current_predicate = None

        # 显示选择支
        add_decision(doc, options)
        return True

    # Predicate: 分支标记
    # [Predicate(references="1")] 或 [Predicate(references="1;2")]
    m = re.search(r'Predicate\(references\s*=\s*"([^"]+)"', line, re.IGNORECASE)
    if m:
        references = m.group(1).strip()

        if assembler:
            assembler.current_predicate = references
            # 显示分支标题
            if assembler.current_decision_options:
                add_predicate_header(doc, references, assembler.current_decision_options)

        return True

    # animtext with <p=1> and <p=2>
    m = re.search(r'<p=1>([^<\n]+)<p=2>([^<\n]+)', line)
    if m:
        # treat as scene title + timestamp
        title = m.group(1).strip()
        time = m.group(2).strip()
        add_scene_title(doc, title)
        add_scene_timestamp(doc, time)
        return True

    # Subtitle(text="...") -> produce as centered short narration
    m = re.search(r'Subtitle\(text\s*=\s*"([^"]+)"', line)
    if m:
        add_narration(doc, m.group(1))
        return True

    # [name="角色"]对话
    m = re.match(r'\[name="([^"]+)"\](.*)', line)
    if m:
        character = m.group(1).strip()
        text = m.group(2).strip()
        if not text:
            text = ''
        add_dialogue(doc, character, text)
        return True

    # lines like name="阿米娅"]Mantra女士？怎么了？ (no bracket start)
    m = re.search(r'name\s*=\s*"([^"]+)"\]\s*(.*)', line)
    if m:
        add_dialogue(doc, m.group(1).strip(), m.group(2).strip())
        return True

    # sound / music related directives
    if line.startswith('[') and ('PlaySound' in line or 'PlayMusic' in line or 'StopSound' in line or 'stopmusic' in line or 'playsound' in line.lower() or 'StopMusic' in line):
        # try extract key
        m = re.search(r'key\s*=\s*"?([^",\)\]]+)"?', line)
        if m:
            key = m.group(1)
            add_sound_effect(doc, key)
        else:
            # 没有 key= 的情况：若指令名在跳过列表中，视为已识别但不输出
            cmd_m = re.match(r'\[?\s*([A-Za-z_][A-Za-z0-9_]*)', line)
            cmd = cmd_m.group(1) if cmd_m else ''
            cmd_lower = cmd.lower()
            SKIP_DIRECTIVES_LOWER = {
                'stopsound', 'stopmusic', 'soundvolume', 'blocker', 'delay',
                'background', 'image', 'imagetween', 'curtain', 'camerashake', 'cameraeffect',
                'focusout', 'bgeffect', 'charslot', 'dialog', 'subtitle', 'animtextclean',
                'animtext', 'playsound', 'playmusic'
            }
            if cmd_lower in SKIP_DIRECTIVES_LOWER:
                # 识别但不写入
                pass
            else:
                # fallback: 原样尝试解析并写入
                add_sound_effect(doc, line.strip('[]'))
        return True

    # Subtitle-like short quoted lines in plain text like: "留下。"
    if re.match(r'^[""].+[""]$', line):
        add_narration(doc, line)
        return True

    # scene separator markers
    if line == '#' or line.lower().startswith('[dialog]') or line.lower().startswith('[charslot]') or line.lower().startswith('[background]'):
        # 明确忽略这些结构化/控制指令（不写入 docx）
        return False

    # lines with angle-bracketed stage directions e.g. <雷声>
    if line.startswith('<') and line.endswith('>'):
        add_sound_effect(doc, line.strip('<>'))
        return True

    # 未识别的行：不写入 docx（按用户要求）
    return False


def parse_lines(lines, outpath, verbose: bool = False):
    """向后兼容的包装：使用 DocumentAssembler 解析并保存。

    Returns list of skipped lines.
    """
    asm = DocumentAssembler()
    skipped = asm.parse_lines(lines)
    asm.save(outpath)
    if verbose:
        skipped_path = outpath + '.skipped.txt'
        try:
            with open(skipped_path, 'w', encoding='utf-8') as sf:
                for l in skipped:
                    sf.write(l + '\n')
            print(f'注意: 有 {len(skipped)} 行被跳过，详情保存在: {skipped_path}')
        except Exception as e:
            print('写入 skipped file 失败:', e)
    return skipped


def parse_file(inpath: str, outpath: str, verbose: bool = False):
    """Parse a file at `inpath` and write output to `outpath`.

    This is an importable function; pass `verbose=True` to write a `.skipped.txt`.
    """
    with open(inpath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    return parse_lines(lines, outpath, verbose=verbose)


def parse_text(text: str, outpath: str, verbose: bool = False):
    """Parse a text blob and write to `outpath`.

    Returns the skipped lines list.
    """
    lines = text.splitlines()
    return parse_lines(lines, outpath, verbose=verbose)


def main():
    if len(sys.argv) < 2:
        print('Usage: python parse_test_to_docx.py [input.txt] [output.docx]')
        sys.exit(1)

    inpath = sys.argv[1]
    outpath = sys.argv[2] if len(sys.argv) >= 3 else 'formatted_test.docx'
    verbose = '--verbose' in sys.argv or '-v' in sys.argv
    try:
        parse_file(inpath, outpath, verbose=verbose)
        print(f'已生成: {outpath}')
    except Exception as e:
        print('解析或保存出错:', e)


if __name__ == '__main__':
    main()
